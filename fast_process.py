import os
import sys
from docx import Document
from docx.oxml.shared import qn
import re
import sqlite3
from typing import Any

db_dir = "books/collections.db"
desc_dir = "books/descriptions"
books_dir = "books"


FIELD_MAP = {
    "Autore": "author",
    "Autore secondario": "second_author", 
    "Titolo": "title",
    "Pubblicazione": "publication",
    "Dimensioni": "dimensions",
    "Peso": "weight",
    "Spessore dei fogli": "thickness",
    "Collocazione": "location",
    "Segnatura": "signature",
    "Impronta": "imprint",
    "Disposizione del testo": "text_layout",
    "Righe": "lines",
    "Richiami": "requests",
    "Legatura": "binding",
    "Lingua": "language_info",
    "Nomi significativi": "significant_names",
    "Stato di conservazione": "condition_info",
    "Decorazione": "decoration",
    "Descrizione fisica": "physical_description",
    "": None  # fallback
}

CINQUECENTINE_COLLECTION_ID = 4
INCUNABOLI_COLLECTION_ID = 3

def show_menu():
    print("\n" + "=" * 60)
    print("Book Agent - What would you like to do?")
    print("=" * 60)
    print("\n1. Process ALL books in 'books/' folder (generate manifests)")
    print("2. Process ONE specific book folder (generate manifest)")
    print("3. Update database from Word documents")
    print("4. Process books + Update database (complete workflow)")
    print("5. List all book folders")
    print("6. Exit")
    print()

def list_books():
    """Show all book folders"""
    
    if not os.path.exists(books_dir):
        print(f"\n❌ '{books_dir}' folder doesn't exist!")
        print(f"   Current directory: {os.getcwd()}")
        return
    
    folders = [d for d in os.listdir(books_dir) 
               if os.path.isdir(os.path.join(books_dir, d)) and d != "descriptions"]
    
    if not folders:
        print(f"\n📁 No book folders found in '{books_dir}'")
        return
    
    print(f"\n📚 Found {len(folders)} book folder(s) in '{books_dir}':")
    for i, folder in enumerate(folders, 1):
        path = os.path.join(books_dir, folder)
        files = os.listdir(path)
        jpg_count = len([f for f in files if f.endswith('.jpg')])
        has_manifest = 'manifest.json' in files
        has_viewer = any(f.startswith('Viewer') and f.endswith('.js') for f in files)
        
        status = "✅ PROCESSED" if (has_manifest and has_viewer) else "⏳ NOT PROCESSED"
        print(f"  {i}. {folder:<20} {status}  ({jpg_count} images)")

def determine_collection_from_filename(filename):
    """Determine which collection this file belongs to based on the number"""
    # Extract number from filename
    match = re.search(r'Scheda descrittiva_([A-Za-z0-9()]+)_VERIFICATA', filename)
    if not match:
        return None, None
    
    file_number = match.group(1).upper()
    
    # Cinquecentine: 5A01, 5B13, etc. (starts with 5)
    if file_number.startswith('5'):
        return CINQUECENTINE_COLLECTION_ID, "cinquecentine"
    
    # Incunaboli: 4B1, 4C2, etc. (starts with 4)  
    elif file_number.startswith('4'):
        return INCUNABOLI_COLLECTION_ID, "incunaboli"
    
    return None, None

def extract_hyperlinks_from_paragraph(paragraph):
    """Extract paragraph text while preserving hyperlinks as HTML.
    
    This handles multiple Word hyperlink structures:
    1. Hyperlinks as run parents (standard)
    2. Hyperlinks in paragraph XML (some Word versions)
    """
    parts = []
    
    # Method 1: Check each run's parent for hyperlink
    for run in paragraph.runs:
        run_parent = run._element.getparent()
        if run_parent is not None and run_parent.tag is not None and run_parent.tag.endswith('hyperlink'):
            rel_id = run_parent.get(qn('r:id'))
            hyperlink_url = None
            if rel_id and rel_id in paragraph.part.rels:
                hyperlink_url = paragraph.part.rels[rel_id].target_ref

            run_text = run.text or ''
            if hyperlink_url and run_text.strip():
                parts.append(f'<a href="{hyperlink_url}" target="_blank">{run_text}</a>')
            else:
                parts.append(run_text)
        else:
            parts.append(run.text or '')

    joined = ''.join(parts)

    # Method 2: If no hyperlinks found in runs, check paragraph XML directly
    if '<a href=' not in joined:
        try:
            # Get paragraph XML
            para_xml = paragraph._element.xml
            if isinstance(para_xml, bytes):
                para_xml = para_xml.decode('utf-8')
            else:
                para_xml = str(para_xml)
            
            # Check if there are hyperlinks in the XML
            if '<w:hyperlink' in para_xml or 'hyperlink' in para_xml.lower():
                # Extract hyperlink elements with their rIds
                import re
                
                # Find all w:hyperlink elements with their r:id
                hyperlink_pattern = r'<w:hyperlink[^>]*r:id="([^"]+)"[^>]*>(.*?)</w:hyperlink>'
                hyperlink_matches = re.findall(hyperlink_pattern, para_xml, re.DOTALL)
                
                if hyperlink_matches:
                    # We found hyperlinks in the XML, need to reconstruct text with links
                    result_parts = []
                    last_pos = 0
                    
                    # Get plain text to know the structure
                    plain_text = paragraph.text
                    
                    for rel_id, hyperlink_content in hyperlink_matches:
                        # Extract the text from the hyperlink content
                        text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', hyperlink_content)
                        link_text = ''.join(text_matches)
                        
                        # Try to get the URL from relationships
                        try:
                            if rel_id in paragraph.part.rels:
                                url = paragraph.part.rels[rel_id].target_ref
                                
                                # Find where this text appears in the plain text
                                if link_text in plain_text[last_pos:]:
                                    pos = plain_text.find(link_text, last_pos)
                                    # Add text before the link
                                    if pos > last_pos:
                                        result_parts.append(plain_text[last_pos:pos])
                                    # Add the link
                                    result_parts.append(f'<a href="{url}" target="_blank">{link_text}</a>')
                                    last_pos = pos + len(link_text)
                        except:
                            # If we can't resolve the URL, just add the text
                            result_parts.append(link_text)
                    
                    # Add any remaining text
                    if last_pos < len(plain_text):
                        result_parts.append(plain_text[last_pos:])
                    
                    joined = ''.join(result_parts)
        
        except Exception as e:
            # If XML parsing fails, fall back to plain text
            pass

    # If there were no hyperlinks, fall back to the simple paragraph text
    if '<a href=' not in joined:
        return (paragraph.text or '').strip()

    return joined

def extract_data_from_docx(docx_path):
    doc = Document(docx_path)
    
    # Extract text with hyperlinks preserved
    full_text_with_links = ""
    paragraph_texts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Try to extract with hyperlinks
            text_with_links = extract_hyperlinks_from_paragraph(paragraph)
            paragraph_texts.append(text_with_links)
            full_text_with_links += text_with_links + "\n"
    
    # Also get plain text for regex matching (fallback)
    plain_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    
    data = {}
    for key, db_col in FIELD_MAP.items():
        if not db_col:
            continue
        
        # First try to find the field in text with links
        match = re.search(rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zÀ-ÿ ]+?:|\Z)", full_text_with_links, re.DOTALL | re.IGNORECASE)
        
        # If not found, try plain text
        if not match:
            match = re.search(rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zÀ-ÿ ]+?:|\Z)", plain_text, re.DOTALL | re.IGNORECASE)
        
        if match:
            value = match.group(1).strip()
            # Clean up whitespace but preserve HTML tags
            value = re.sub(r'\s+', ' ', value)
            # Clean up any malformed HTML
            value = re.sub(r'<a\s+href="([^"]*)"[^>]*>\s*</a>', '', value)  # Remove empty links
            data[db_col] = value
            
    return data

def normalize_book_number(number):
    """Normalize book number to handle different formats like 5A01 vs 5A1"""
    if not number:
        return ""
    
    # Remove any extra spaces and convert to uppercase
    number = str(number).strip().upper()
    
    # Handle patterns like 5A01 -> 5A1 (remove leading zero in the number part)
    match = re.match(r'^(\d+)([A-Z]+)(\d+)$', number)
    if match:
        prefix, letter, num = match.groups()
        # Remove leading zeros from the number part
        normalized_num = str(int(num))
        return f"{prefix}{letter}{normalized_num}"
    
    return number

def find_or_create_book(cursor, file_number, collection_id, author=None):
    """Find a book in the specified collection that matches the file number, or create it if not found"""
    
    # Normalize the file number
    normalized_file_number = normalize_book_number(file_number)
    
    # Try to find exact match first
    cursor.execute(
        "SELECT book_id, number FROM books WHERE collection_id = ? AND UPPER(number) = ?", 
        (collection_id, file_number.upper())
    )
    result = cursor.fetchone()
    if result:
        return result, "found"
    
    # Try normalized match
    cursor.execute(
        "SELECT book_id, number FROM books WHERE collection_id = ?", 
        (collection_id,)
    )
    books = cursor.fetchall()
    
    for book_id, db_number in books:
        if normalize_book_number(db_number) == normalized_file_number:
            return (book_id, db_number), "found"
    
    # Book not found - create it
    print(f"   📝 Creating new book: {file_number} in database")
    try:
        insert_data = {
            'collection_id': collection_id,
            'number': file_number,
            'author': author or 'Unknown'
        }
        
        columns = list(insert_data.keys())
        placeholders = ['?' for _ in columns]
        sql = f"INSERT INTO books ({', '.join(columns)}) VALUES ({', '.join(placeholders)})"
        
        cursor.execute(sql, list(insert_data.values()))
        book_id = cursor.lastrowid
        
        return (book_id, file_number), "created"
    except Exception as e:
        return None, f"error: {str(e)}"

def insert_or_update_description(cursor, book_id, collection_id, book_number, data):
    """Insert or update book description in the book_descriptions table"""
    
    # Check if description already exists
    cursor.execute(
        "SELECT description_id FROM book_descriptions WHERE book_id = ? AND language = 'it'",
        (book_id,)
    )
    existing = cursor.fetchone()
    
    # Prepare full data including required fields
    full_data = {
        'book_id': book_id,
        'collection_id': collection_id,
        'number': book_number,
        'language': 'it',
        **data  # Add extracted description data
    }
    
    if existing:
        # Update existing description
        description_id = existing[0]
        set_clauses = []
        values = []
        
        for key, value in data.items():
            # Ensure we're storing the full text with HTML tags
            if value:
                set_clauses.append(f"{key} = ?")
                values.append(str(value))  # Convert to string to preserve HTML
        
        if set_clauses:
            values.append(description_id)
            sql = f"UPDATE book_descriptions SET {', '.join(set_clauses)} WHERE description_id = ?"
            cursor.execute(sql, values)
            return "updated"
        return "no_changes"
    else:
        # Insert new description - make sure all values are strings
        clean_data = {k: str(v) if v else None for k, v in full_data.items()}
        columns = list(clean_data.keys())
        placeholders = ['?' for _ in columns]
        
        sql = f"INSERT INTO book_descriptions ({', '.join(columns)}) VALUES ({', '.join(placeholders)})"
        cursor.execute(sql, list(clean_data.values()))
        return "inserted"

def update_db_from_docx():
    """Update database from Word documents"""
    print("\n" + "=" * 60)
    print("Updating Database from Word Documents")
    print("=" * 60)
    
    if not os.path.exists(db_dir):
        print(f"\n❌ Database not found: {db_dir}")
        return
    
    if not os.path.exists(desc_dir):
        print(f"\n❌ Descriptions folder not found: {desc_dir}")
        return
    
    conn = sqlite3.connect(db_dir)
    cursor = conn.cursor()

    # Check if book_descriptions table exists
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='book_descriptions'")
    if not cursor.fetchone():
        print("\n❌ book_descriptions table does not exist. Please run migration script first.")
        conn.close()
        return

    # Get current counts
    cursor.execute(
        "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?", 
        (CINQUECENTINE_COLLECTION_ID,)
    )
    cinquecentine_books = cursor.fetchone()[0]
    
    cursor.execute(
        "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?", 
        (INCUNABOLI_COLLECTION_ID,)
    )
    incunaboli_books = cursor.fetchone()[0]
    
    print(f"\n📚 Current database state:")
    print(f"   Cinquecentine: {cinquecentine_books} descriptions")
    print(f"   Incunaboli: {incunaboli_books} descriptions")

    processed_count = 0
    updated_count = 0
    inserted_count = 0
    books_created_count = 0
    not_found_count = 0

    docx_files = [f for f in os.listdir(desc_dir) if f.endswith((".docx", ".doc"))]
    print(f"\n📄 Found {len(docx_files)} Word document(s) to process")
    print()

    for filename in docx_files:
        # Extract number and determine collection
        match = re.search(r'Scheda descrittiva_([A-Za-z0-9()]+)_VERIFICATA', filename)
        if not match:
            print(f"❌ Could not extract number from: {filename}")
            continue

        file_number = match.group(1).upper()
        collection_id, collection_name = determine_collection_from_filename(filename)
        
        if not collection_id:
            print(f"❌ Could not determine collection for: {filename} (number: {file_number})")
            continue
        
        docx_path = os.path.join(desc_dir, filename)
        
        print(f"🔍 Processing: {filename}")
        print(f"   📂 Collection: {collection_name} (ID: {collection_id})")
        print(f"   📖 Book number: {file_number}")
        
        # Extract data first to get author
        try:
            data = extract_data_from_docx(docx_path)
        except Exception as e:
            print(f"   ❌ Error processing: {e}")
            continue

        if not data:
            print(f"   ⚠️  No metadata extracted")
            continue
        
        # Find or create book in database
        author = data.get('author', 'Unknown')
        book_result = find_or_create_book(cursor, file_number, collection_id, author)
        
        if book_result[1].startswith("error"):
            print(f"   ❌ Error: {book_result[1]}")
            not_found_count += 1
            continue
        
        book_match, status = book_result
        
        if book_match is None:
            print(f"   ❌ Could not find or create book")
            not_found_count += 1
            continue
        
        book_id, db_number = book_match
        
        if status == "created":
            print(f"   ✨ Created new book: {db_number} (ID: {book_id})")
            books_created_count += 1
        else:
            print(f"   ✅ Found book: {db_number} (ID: {book_id})")

        # Insert or update description
        try:
            action = insert_or_update_description(cursor, book_id, collection_id, db_number, data)
            if action == "updated":
                print(f"   ✅ Updated description")
                updated_count += 1
            elif action == "inserted":
                print(f"   ✅ Inserted description")
                inserted_count += 1
            else:
                print(f"   ℹ️  No changes needed")
            
            # Verify and show what was actually stored
            hyperlink_fields = []
            for col, value in data.items():
                if value and '<a href=' in str(value):
                    hyperlink_fields.append(col)
                    # Show a preview
                    preview = str(value)[:150]
                    print(f"   📎 Stored with hyperlinks in '{col}': {preview}...")
            
            if hyperlink_fields:
                # Verify it's actually in the database
                cursor.execute(
                    f"SELECT {', '.join(hyperlink_fields)} FROM book_descriptions WHERE book_id = ?",
                    (book_id,)
                )
                db_values = cursor.fetchone()
                if db_values:
                    for i, field in enumerate(hyperlink_fields):
                        if db_values[i] and '<a href=' in str(db_values[i]):
                            print(f"   ✓ Verified '{field}' contains hyperlinks in DB")
                        else:
                            print(f"   ⚠️  WARNING: '{field}' hyperlinks may not be stored correctly!")
                            print(f"      Stored value: {str(db_values[i])[:100]}")
                            
        except sqlite3.Error as e:
            print(f"   ❌ Database error: {e}")
            import traceback
            traceback.print_exc()
        
        processed_count += 1
        print()

    conn.commit()
    conn.close()
    
    print("=" * 60)
    print("📊 Summary:")
    print("=" * 60)
    print(f"   📄 Files processed: {processed_count}")
    print(f"   ✨ Books created: {books_created_count}")
    print(f"   ✅ Descriptions inserted: {inserted_count}")
    print(f"   🔄 Descriptions updated: {updated_count}")
    print(f"   ❌ Books not found: {not_found_count}")
    print("=" * 60)

def process_all():
    """Process all book folders"""
    from generator import process_book_folder
    
    if not os.path.exists(books_dir):
        print(f"\n❌ '{books_dir}' folder doesn't exist!")
        return
    
    folders = [d for d in os.listdir(books_dir) 
               if os.path.isdir(os.path.join(books_dir, d)) and d != "descriptions"]
    
    if not folders:
        print(f"\n⚠️  No book folders found!")
        return
    
    print(f"\n📚 Processing {len(folders)} book folder(s)...")
    print("=" * 60)
    
    success_count = 0
    error_count = 0
    
    for folder in folders:
        path = os.path.join(books_dir, folder)
        print(f"\n📖 {folder}...")
        
        try:
            process_book_folder(path)
            print(f"   ✅ Success")
            success_count += 1
        except Exception as e:
            print(f"   ❌ Error: {e}")
            error_count += 1
    
    print("\n" + "=" * 60)
    print(f"✅ Successfully processed: {success_count}")
    if error_count > 0:
        print(f"❌ Errors: {error_count}")
    print("=" * 60)

def process_one():
    """Process one specific book folder"""
    from generator import process_book_folder
    
    # Show available folders
    if os.path.exists(books_dir):
        folders = [d for d in os.listdir(books_dir) 
                   if os.path.isdir(os.path.join(books_dir, d)) and d != "descriptions"]
        
        if folders:
            print("\nAvailable book folders:")
            for folder in folders:
                print(f"  - {folder}")
    
    print("\nEnter the book folder name (e.g., '5d23'):")
    folder_name = input("> ").strip()
    
    if not folder_name:
        print("❌ No folder name entered")
        return
    
    # Try both with and without 'books/' prefix
    if os.path.exists(folder_name):
        path = folder_name
    elif os.path.exists(os.path.join(books_dir, folder_name)):
        path = os.path.join(books_dir, folder_name)
    else:
        print(f"\n❌ Folder not found: {folder_name}")
        print(f"   Tried: {folder_name}")
        print(f"   Tried: {os.path.join(books_dir, folder_name)}")
        return
    
    print(f"\n📖 Processing: {path}")
    print("=" * 60)
    
    try:
        process_book_folder(path)
        print("\n" + "=" * 60)
        print("✅ Success!")
        print("=" * 60)
        print(f"\nCheck {path}/ for:")
        print("  - manifest.json")
        print("  - Viewer*.js")
    except Exception as e:
        print("\n" + "=" * 60)
        print("❌ Error!")
        print("=" * 60)
        print(f"{e}")
        import traceback
        traceback.print_exc()

def process_complete_workflow():
    """Process books and update database - complete workflow"""
    print("\n" + "=" * 60)
    print("Complete Workflow: Process Books + Update Database")
    print("=" * 60)
    
    print("\n🔹 Step 1: Processing book folders...")
    process_all()
    
    print("\n🔹 Step 2: Updating database...")
    update_db_from_docx()
    
    print("\n" + "=" * 60)
    print("✅ Complete workflow finished!")
    print("=" * 60)

def main():
    print("\n" + "=" * 60)
    print("Book Agent - Interactive Book Processor")
    print("=" * 60)
    print(f"\nCurrent directory: {os.getcwd()}")
    
    # Check if we're in the right place
    if not os.path.exists("generator.py"):
        print("\n⚠️  Warning: Can't find generator.py")
        print("   Make sure you're in the book_agent directory")
        print(f"\n   Run: cd ~/projects/book_agent")
        return
    
    while True:
        show_menu()
        choice = input("Your choice (1-6): ").strip()
        
        if choice == "1":
            process_all()
        elif choice == "2":
            process_one()
        elif choice == "3":
            update_db_from_docx()
        elif choice == "4":
            process_complete_workflow()
        elif choice == "5":
            list_books()
        elif choice == "6":
            print("\n👋 Goodbye!")
            break
        else:
            print(f"\n❌ Invalid choice: {choice}")
        
        input("\nPress Enter to continue...")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n👋 Goodbye!")
        sys.exit(0)