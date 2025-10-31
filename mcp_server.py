import asyncio
import os
import sqlite3
import re
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.oxml.shared import qn

from mcp.server import Server
from mcp.server.models import InitializationOptions
import mcp.types as types

# Import your existing modules
from watcher import start_watching
from generator import process_book_folder

# Initialize server
server = Server("book-agent")

# Configuration - Your paths
SCRIPT_DIR = Path(__file__).resolve().parent
WATCH_PATH = str(SCRIPT_DIR / "books")
DOCX_FOLDER = str(SCRIPT_DIR / "books" / "descriptions")
DB_PATH = str(SCRIPT_DIR / "books" / "collections.db")

# Collection IDs
CINQUECENTINE_COLLECTION_ID = 4
INCUNABOLI_COLLECTION_ID = 3

# Field mapping for Italian to DB columns
FIELD_MAP = {
    "Autore": "author",
    "Autore secondario": "second_author",
    "Titolo": "title",
    "Pubblicazione": "publication",
    "Dimensioni": "dimensions",
    "Peso": "weight",
    "Spessore dei fogli": "thickness",
    "Collocazione": "location",
    "Segnatura": "signature",
    "Impronta": "imprint",
    "Disposizione del testo": "text_layout",
    "Righe": "lines",
    "Richiami": "requests",
    "Legatura": "binding",
    "Lingua": "language_info",
    "Nomi significativi": "significant_names",
    "Stato di conservazione": "condition_info",
    "Decorazione": "decoration",
    "Descrizione fisica": "physical_description",
}

# Store paths in a mutable dict to avoid global declaration issues
CONFIG = {
    "watch_path": WATCH_PATH,
    "docx_folder": DOCX_FOLDER,
    "db_path": DB_PATH
}


# ========== HELPER FUNCTIONS ==========

def normalize_book_number(number: str) -> str:
    """Normalize book number (e.g., 5A01 -> 5A1)"""
    if not number:
        return ""
    number = str(number).strip().upper()
    match = re.match(r'^(\d+)([A-Z]+)(\d+)$', number)
    if match:
        prefix, letter, num = match.groups()
        normalized_num = str(int(num))
        return f"{prefix}{letter}{normalized_num}"
    return number


def determine_collection_from_filename(filename: str):
    """Determine collection ID from filename"""
    match = re.search(r'Scheda descrittiva_([A-Za-z0-9()]+)_VERIFICATA', filename)
    if not match:
        return None, None
    
    file_number = match.group(1).upper()
    
    if file_number.startswith('5'):
        return CINQUECENTINE_COLLECTION_ID, "cinquecentine"
    elif file_number.startswith('4'):
        return INCUNABOLI_COLLECTION_ID, "incunaboli"
    
    return None, None


def extract_hyperlinks_from_paragraph(paragraph):
    """Extract paragraph text preserving hyperlinks as HTML"""
    parts = []
    for run in paragraph.runs:
        run_parent = run._element.getparent()
        if run_parent is not None and run_parent.tag and run_parent.tag.endswith('hyperlink'):
            rel_id = run_parent.get(qn('r:id'))
            hyperlink_url = None
            if rel_id and rel_id in paragraph.part.rels:
                hyperlink_url = paragraph.part.rels[rel_id].target_ref
            
            run_text = run.text or ''
            if hyperlink_url and run_text.strip():
                parts.append(f'<a href="{hyperlink_url}" target="_blank">{run_text}</a>')
            else:
                parts.append(run_text)
        else:
            parts.append(run.text or '')
    
    joined = ''.join(parts)
    if '<a href=' not in joined:
        return (paragraph.text or '').strip()
    return joined


def extract_data_from_docx(docx_path: str) -> Dict[str, Any]:
    """Extract structured data from DOCX file"""
    doc = Document(docx_path)
    
    # Extract text with hyperlinks
    full_text_with_links = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_with_links = extract_hyperlinks_from_paragraph(paragraph)
            full_text_with_links += text_with_links + "\n"
    
    # Plain text fallback
    plain_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    
    data = {}
    for key, db_col in FIELD_MAP.items():
        if not db_col:
            continue
        
        # Try with links first
        match = re.search(
            rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zà-ÿ ]+?:|\Z)",
            full_text_with_links,
            re.DOTALL | re.IGNORECASE
        )
        
        # Fallback to plain text
        if not match:
            match = re.search(
                rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zà-ÿ ]+?:|\Z)",
                plain_text,
                re.DOTALL | re.IGNORECASE
            )
        
        if match:
            value = match.group(1).strip()
            value = re.sub(r'\s+', ' ', value)
            value = re.sub(r'<a\s+href="([^"]*)"[^>]*>\s*</a>', '', value)
            data[db_col] = value
    
    return data


def find_or_create_book(cursor, file_number: str, collection_id: int, author: str = None):
    """Find matching book in database, or create it if not found"""
    normalized_file_number = normalize_book_number(file_number)
    
    # Try exact match
    cursor.execute(
        "SELECT book_id, number FROM books WHERE collection_id = ? AND UPPER(number) = ?",
        (collection_id, file_number.upper())
    )
    result = cursor.fetchone()
    if result:
        return result, "found"
    
    # Try normalized match
    cursor.execute(
        "SELECT book_id, number FROM books WHERE collection_id = ?",
        (collection_id,)
    )
    books = cursor.fetchall()
    
    for book_id, db_number in books:
        if normalize_book_number(db_number) == normalized_file_number:
            return (book_id, db_number), "found"
    
    # Book not found - create it
    try:
        # Insert new book
        insert_data = {
            'collection_id': collection_id,
            'number': file_number,
            'author': author or 'Unknown'
        }
        
        columns = list(insert_data.keys())
        placeholders = ['?' for _ in columns]
        sql = f"INSERT INTO books ({', '.join(columns)}) VALUES ({', '.join(placeholders)})"
        
        cursor.execute(sql, list(insert_data.values()))
        book_id = cursor.lastrowid
        
        return (book_id, file_number), "created"
    except Exception as e:
        return None, f"error: {str(e)}"


def insert_or_update_description(cursor, book_id: int, collection_id: int, 
                                 book_number: str, data: Dict[str, Any]) -> str:
    """Insert or update book description"""
    cursor.execute(
        "SELECT description_id FROM book_descriptions WHERE book_id = ? AND language = 'it'",
        (book_id,)
    )
    existing = cursor.fetchone()
    
    full_data = {
        'book_id': book_id,
        'collection_id': collection_id,
        'number': book_number,
        'language': 'it',
        **data
    }
    
    if existing:
        description_id = existing[0]
        set_clauses = []
        values = []
        
        for key, value in data.items():
            set_clauses.append(f"{key} = ?")
            values.append(value)
        
        if set_clauses:
            values.append(description_id)
            sql = f"UPDATE book_descriptions SET {', '.join(set_clauses)} WHERE description_id = ?"
            cursor.execute(sql, values)
            return "updated"
    else:
        columns = list(full_data.keys())
        placeholders = ['?' for _ in columns]
        
        sql = f"INSERT INTO book_descriptions ({', '.join(columns)}) VALUES ({', '.join(placeholders)})"
        cursor.execute(sql, list(full_data.values()))
        return "inserted"


def _update_descriptions_sync(folder: str, db: str) -> dict:
    """Synchronous version of update_descriptions_from_docx"""
    conn = sqlite3.connect(db)
    cursor = conn.cursor()
    
    # Check table exists
    cursor.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='book_descriptions'"
    )
    if not cursor.fetchone():
        conn.close()
        return {
            "error": "book_descriptions table does not exist",
            "message": "Please run migration script first"
        }
    
    # Get current counts
    cursor.execute(
        "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?",
        (CINQUECENTINE_COLLECTION_ID,)
    )
    cinquecentine_books = cursor.fetchone()[0]
    
    cursor.execute(
        "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?",
        (INCUNABOLI_COLLECTION_ID,)
    )
    incunaboli_books = cursor.fetchone()[0]
    
    # Statistics
    stats = {
        "processed": 0,
        "inserted": 0,
        "updated": 0,
        "not_found": 0,
        "books_created": 0,
        "errors": [],
        "details": []
    }
    
    # Process each DOCX file
    for filename in os.listdir(folder):
        if not filename.endswith((".docx", ".doc")):
            continue
        
        # Extract number and collection
        match = re.search(r'Scheda descrittiva_([A-Za-z0-9()]+)_VERIFICATA', filename)
        if not match:
            stats["errors"].append(f"Could not parse: {filename}")
            continue
        
        file_number = match.group(1).upper()
        collection_id, collection_name = determine_collection_from_filename(filename)
        
        if not collection_id:
            stats["errors"].append(f"Unknown collection: {filename}")
            continue
        
        docx_path = os.path.join(folder, filename)
        
        # Extract data first to get author name
        try:
            data = extract_data_from_docx(docx_path)
            if not data:
                stats["errors"].append(f"No data extracted: {filename}")
                continue
        except Exception as e:
            stats["errors"].append(f"Error processing {filename}: {str(e)}")
            continue
        
        # Find or create book
        author = data.get('author', 'Unknown')
        book_result = find_or_create_book(cursor, file_number, collection_id, author)
        
        if book_result[1].startswith("error"):
            stats["errors"].append(f"Error with book {file_number}: {book_result[1]}")
            continue
        
        book_match, status = book_result
        
        if book_match is None:
            stats["not_found"] += 1
            stats["errors"].append(f"Could not find or create book: {file_number} in {collection_name}")
            continue
        
        book_id, db_number = book_match
        
        # Track if we created a new book
        if status == "created":
            stats["books_created"] += 1
        
        # Insert or update description
        try:
            action = insert_or_update_description(cursor, book_id, collection_id, db_number, data)
            
            detail = {
                "filename": filename,
                "book_number": db_number,
                "collection": collection_name,
                "action": action,
                "book_status": status
            }
            
            # Check for hyperlinks
            for col, value in list(data.items())[:2]:
                if '<a href=' in str(value):
                    detail["has_hyperlinks"] = True
                    break
            
            stats["details"].append(detail)
            
            if action == "updated":
                stats["updated"] += 1
            else:
                stats["inserted"] += 1
            
            stats["processed"] += 1
            
        except Exception as e:
            stats["errors"].append(f"Database error for book {db_number}: {str(e)}")
    
    conn.commit()
    conn.close()
    
    return {
        "success": True,
        "statistics": stats,
        "message": f"Processed {stats['processed']} files, created {stats['books_created']} new books",
        "summary": {
            "cinquecentine_books": cinquecentine_books,
            "incunaboli_books": incunaboli_books,
            "files_processed": stats["processed"],
            "descriptions_inserted": stats["inserted"],
            "descriptions_updated": stats["updated"],
            "books_created": stats["books_created"],
            "books_not_found": stats["not_found"]
        }
    }


# ========== MCP TOOLS ==========

@server.list_tools()
async def handle_list_tools() -> list[types.Tool]:
    """List available tools"""
    return [
        types.Tool(
            name="update_descriptions_from_docx",
            description="Update book descriptions in database from DOCX files. Reads Word documents from the descriptions folder and updates the book_descriptions table.",
            inputSchema={
                "type": "object",
                "properties": {
                    "docx_folder": {
                        "type": "string",
                        "description": f"Path to folder containing DOCX files (default: {CONFIG['docx_folder']})"
                    },
                    "db_path": {
                        "type": "string",
                        "description": f"Path to database file (default: {CONFIG['db_path']})"
                    }
                }
            }
        ),
        types.Tool(
            name="process_new_books",
            description="Process new book folders and generate IIIF manifests and React viewer components",
            inputSchema={
                "type": "object",
                "properties": {}
            }
        ),
        types.Tool(
            name="get_database_stats",
            description="Get statistics about the database including book counts by collection",
            inputSchema={
                "type": "object",
                "properties": {
                    "db_path": {
                        "type": "string",
                        "description": f"Path to database file (default: {CONFIG['db_path']})"
                    }
                }
            }
        ),
        types.Tool(
            name="configure_paths",
            description="Update configuration paths for watch folder, database, and DOCX folder",
            inputSchema={
                "type": "object",
                "properties": {
                    "watch_path": {
                        "type": "string",
                        "description": "Path to watch for new books"
                    },
                    "db_path": {
                        "type": "string",
                        "description": "Path to SQLite database"
                    },
                    "docx_folder": {
                        "type": "string",
                        "description": "Path to folder with DOCX description files"
                    }
                }
            }
        )
    ]


@server.call_tool()
async def handle_call_tool(
    name: str, arguments: dict | None
) -> list[types.TextContent | types.ImageContent | types.EmbeddedResource]:
    """Handle tool execution"""
    
    if arguments is None:
        arguments = {}
    
    if name == "update_descriptions_from_docx":
        folder = arguments.get("docx_folder", CONFIG["docx_folder"])
        db = arguments.get("db_path", CONFIG["db_path"])
        
        if not os.path.exists(folder):
            return [types.TextContent(
                type="text",
                text=f"❌ Error: DOCX folder not found: {folder}\n\nCurrent config:\n  DOCX folder: {CONFIG['docx_folder']}\n  Database: {CONFIG['db_path']}"
            )]
        
        if not os.path.exists(db):
            return [types.TextContent(
                type="text",
                text=f"❌ Error: Database not found: {db}\n\nCurrent config:\n  DOCX folder: {CONFIG['docx_folder']}\n  Database: {CONFIG['db_path']}"
            )]
        
        # Run in thread pool
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, _update_descriptions_sync, folder, db)
        
        import json
        return [types.TextContent(
            type="text",
            text=json.dumps(result, indent=2)
        )]
    
    elif name == "process_new_books":
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, process_book_folder, CONFIG["watch_path"])
        
        import json
        return [types.TextContent(
            type="text",
            text=json.dumps({"success": True, "message": "Books processed successfully"}, indent=2)
        )]
    
    elif name == "get_database_stats":
        db = arguments.get("db_path", CONFIG["db_path"])
        
        if not os.path.exists(db):
            return [types.TextContent(
                type="text",
                text=f"❌ Error: Database not found: {db}"
            )]
        
        try:
            conn = sqlite3.connect(db)
            cursor = conn.cursor()
            
            cursor.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='book_descriptions'"
            )
            has_descriptions = cursor.fetchone() is not None
            
            stats = {
                "database_path": db,
                "has_descriptions_table": has_descriptions
            }
            
            if has_descriptions:
                cursor.execute(
                    "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?",
                    (CINQUECENTINE_COLLECTION_ID,)
                )
                stats["cinquecentine_descriptions"] = cursor.fetchone()[0]
                
                cursor.execute(
                    "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?",
                    (INCUNABOLI_COLLECTION_ID,)
                )
                stats["incunaboli_descriptions"] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM books WHERE collection_id = ?", (CINQUECENTINE_COLLECTION_ID,))
            stats["cinquecentine_books"] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM books WHERE collection_id = ?", (INCUNABOLI_COLLECTION_ID,))
            stats["incunaboli_books"] = cursor.fetchone()[0]
            
            conn.close()
            
            import json
            return [types.TextContent(
                type="text",
                text=json.dumps({"success": True, "statistics": stats}, indent=2)
            )]
        
        except Exception as e:
            return [types.TextContent(
                type="text",
                text=f"❌ Error: {str(e)}"
            )]
    
    elif name == "configure_paths":
        if "watch_path" in arguments:
            CONFIG["watch_path"] = arguments["watch_path"]
        if "db_path" in arguments:
            CONFIG["db_path"] = arguments["db_path"]
        if "docx_folder" in arguments:
            CONFIG["docx_folder"] = arguments["docx_folder"]
        
        import json
        return [types.TextContent(
            type="text",
            text=json.dumps({
                "success": True,
                "message": "Configuration updated",
                "config": CONFIG
            }, indent=2)
        )]
    
    return [types.TextContent(
        type="text",
        text=f"❌ Unknown tool: {name}"
    )]


async def main():
    """Run the server"""
    from mcp.server.stdio import stdio_server
    
    async with stdio_server() as (read_stream, write_stream):
        init_options = InitializationOptions(
            server_name="book-agent",
            server_version="1.0.0",
            capabilities=types.ServerCapabilities(
                tools=types.ToolsCapability()
            )
        )
        await server.run(read_stream, write_stream, init_options)


if __name__ == "__main__":
    asyncio.run(main())