#!/usr/bin/env python3
import os
from pathlib import Path
import re
import sqlite3
from typing import Any
from docx import Document
from docx.oxml.shared import qn

# Resolve directories so paths work regardless of current working directory
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent

# Path to your DOCX files (inside scripts folder)
DOCX_FOLDER = str(SCRIPT_DIR / "descriptions")
# Database lives under project public/data
DB_PATH = str(PROJECT_ROOT / "public" / "data" / "collections.db")

# Collection IDs (from the database)
CINQUECENTINE_COLLECTION_ID = 4
INCUNABOLI_COLLECTION_ID = 3

# === MAPPING: Italian field → DB column ===
FIELD_MAP = {
    "Autore": "author",
    "Autore secondario": "second_author", 
    "Titolo": "title",
    "Pubblicazione": "publication",
    "Dimensioni": "dimensions",
    "Peso": "weight",
    "Spessore dei fogli": "thickness",
    "Collocazione": "location",
    "Segnatura": "signature",
    "Impronta": "imprint",
    "Disposizione del testo": "text_layout",
    "Righe": "lines",
    "Richiami": "requests",
    "Legatura": "binding",
    "Lingua": "language_info",
    "Nomi significativi": "significant_names",
    "Stato di conservazione": "condition_info",
    "Decorazione": "decoration",
    "Descrizione fisica": "physical_description",
    "": None  # fallback
}

def normalize_book_number(number):
    """Normalize book number to handle different formats like 5A01 vs 5A1"""
    if not number:
        return ""
    
    # Remove any extra spaces and convert to uppercase
    number = str(number).strip().upper()
    
    # Handle patterns like 5A01 -> 5A1 (remove leading zero in the number part)
    match = re.match(r'^(\d+)([A-Z]+)(\d+)$', number)
    if match:
        prefix, letter, num = match.groups()
        # Remove leading zeros from the number part
        normalized_num = str(int(num))
        return f"{prefix}{letter}{normalized_num}"
    
    return number

def determine_collection_from_filename(filename):
    """Determine which collection this file belongs to based on the number"""
    # Extract number from filename
    match = re.search(r'Scheda descrittiva_([A-Za-z0-9()]+)_VERIFICATA', filename)
    if not match:
        return None, None
    
    file_number = match.group(1).upper()
    
    # Cinquecentine: 5A01, 5B13, etc. (starts with 5)
    if file_number.startswith('5'):
        return CINQUECENTINE_COLLECTION_ID, "cinquecentine"
    
    # Incunaboli: 4B1, 4C2, etc. (starts with 4)  
    elif file_number.startswith('4'):
        return INCUNABOLI_COLLECTION_ID, "incunaboli"
    
    return None, None

def extract_hyperlinks_from_paragraph(paragraph):
    """Extract paragraph text while preserving the exact run order so
    punctuation and link placement remain the same as in Word.

    We iterate over the paragraph.runs in document order. If a run is
    wrapped in a hyperlink element we try to resolve its relationship id
    to a target URL and emit an anchor tag for that run's text. Otherwise
    we emit the run text verbatim. This preserves punctuation that may
    be placed in separate runs after or before the linked text.
    """

    parts = []

    for run in paragraph.runs:
        # Determine whether this run is inside a <w:hyperlink> element
        run_parent = run._element.getparent()
        if run_parent is not None and run_parent.tag is not None and run_parent.tag.endswith('hyperlink'):
            rel_id = run_parent.get(qn('r:id'))
            hyperlink_url = None
            if rel_id and rel_id in paragraph.part.rels:
                hyperlink_url = paragraph.part.rels[rel_id].target_ref

            run_text = run.text or ''
            if hyperlink_url and run_text.strip():
                parts.append(f'<a href="{hyperlink_url}" target="_blank">{run_text}</a>')
            else:
                parts.append(run_text)
        else:
            parts.append(run.text or '')

    joined = ''.join(parts)

    # If there were no hyperlinks, fall back to the simple paragraph text
    if '<a href=' not in joined:
        return (paragraph.text or '').strip()

    return joined

# === HELPER ===
def extract_data_from_docx(docx_path):
    doc = Document(docx_path)
    
    # Extract text with hyperlinks preserved
    full_text_with_links = ""
    paragraph_texts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Try to extract with hyperlinks
            text_with_links = extract_hyperlinks_from_paragraph(paragraph)
            paragraph_texts.append(text_with_links)
            full_text_with_links += text_with_links + "\n"
    
    # Also get plain text for regex matching (fallback)
    plain_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    
    data = {}
    for key, db_col in FIELD_MAP.items():
        if not db_col:
            continue
        
        # First try to find the field in text with links
        match = re.search(rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zÀ-ÿ ]+?:|\Z)", full_text_with_links, re.DOTALL | re.IGNORECASE)
        
        # If not found, try plain text
        if not match:
            match = re.search(rf"{re.escape(key)}:\s*(.+?)(?=\n[A-ZÀ-ÿ][a-zÀ-ÿ ]+?:|\Z)", plain_text, re.DOTALL | re.IGNORECASE)
        
        if match:
            value = match.group(1).strip()
            # Clean up whitespace but preserve HTML tags
            value = re.sub(r'\s+', ' ', value)
            # Clean up any malformed HTML
            value = re.sub(r'<a\s+href="([^"]*)"[^>]*>\s*</a>', '', value)  # Remove empty links
            data[db_col] = value
            
    return data

def find_matching_book(cursor, file_number, collection_id):
    """Find a book in the specified collection that matches the file number"""
    
    # Normalize the file number
    normalized_file_number = normalize_book_number(file_number)
    
    # Try to find exact match first
    cursor.execute(
        "SELECT book_id, number FROM books WHERE collection_id = ? AND UPPER(number) = ?", 
        (collection_id, file_number.upper())
    )
    result = cursor.fetchone()
    if result:
        return result
    
    # Try normalized match
    cursor.execute(
        "SELECT book_id, number FROM books WHERE collection_id = ?", 
        (collection_id,)
    )
    books = cursor.fetchall()
    
    for book_id, db_number in books:
        if normalize_book_number(db_number) == normalized_file_number:
            return (book_id, db_number)
    
    return None

def insert_or_update_description(cursor, book_id, collection_id, book_number, data):
    """Insert or update book description in the book_descriptions table"""
    
    # Check if description already exists
    cursor.execute(
        "SELECT description_id FROM book_descriptions WHERE book_id = ? AND language = 'it'",
        (book_id,)
    )
    existing = cursor.fetchone()
    
    # Prepare full data including required fields
    full_data = {
        'book_id': book_id,
        'collection_id': collection_id,
        'number': book_number,
        'language': 'it',
        **data  # Add extracted description data
    }
    
    if existing:
        # Update existing description
        description_id = existing[0]
        set_clauses = []
        values = []
        
        for key, value in data.items():
            set_clauses.append(f"{key} = ?")
            values.append(value)
        
        if set_clauses:
            values.append(description_id)
            sql = f"UPDATE book_descriptions SET {', '.join(set_clauses)} WHERE description_id = ?"
            cursor.execute(sql, values)
            return "updated"
    else:
        # Insert new description
        columns = list[Any](full_data.keys())
        placeholders = ['?' for _ in columns]
        
        sql = f"INSERT INTO book_descriptions ({', '.join(columns)}) VALUES ({', '.join(placeholders)})"
        cursor.execute(sql, list(full_data.values()))
        return "inserted"

# === MAIN SCRIPT ===
def update_db_from_docx():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    # Check if book_descriptions table exists
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='book_descriptions'")
    if not cursor.fetchone():
        print("❌ book_descriptions table does not exist. Please run migration script first:")
        print("   python scripts/migrate_to_better_schema.py")
        return

    # First, let's see what books we have in both collections
    cursor.execute(
        "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?", 
        (CINQUECENTINE_COLLECTION_ID,)
    )
    cinquecentine_books = cursor.fetchone()[0]
    
    cursor.execute(
        "SELECT COUNT(*) FROM book_descriptions WHERE collection_id = ?", 
        (INCUNABOLI_COLLECTION_ID,)
    )
    incunaboli_books = cursor.fetchone()[0]
    
    print(f"📚 Found {cinquecentine_books} books in cinquecentine collection (ID: {CINQUECENTINE_COLLECTION_ID})")
    print(f"📚 Found {incunaboli_books} books in incunaboli collection (ID: {INCUNABOLI_COLLECTION_ID})")

    processed_count = 0
    updated_count = 0
    inserted_count = 0
    not_found_count = 0

    for filename in os.listdir(DOCX_FOLDER):
        if not filename.endswith((".docx", ".doc")):
            continue

        # Extract number and determine collection
        match = re.search(r'Scheda descrittiva_([A-Za-z0-9()]+)_VERIFICATA', filename)
        if not match:
            print(f"❌ Could not extract number from: {filename}")
            continue

        file_number = match.group(1).upper()  # e.g., "5A01" or "4B1"
        collection_id, collection_name = determine_collection_from_filename(filename)
        
        if not collection_id:
            print(f"❌ Could not determine collection for: {filename} (number: {file_number})")
            continue
        
        docx_path = os.path.join(DOCX_FOLDER, filename)
        
        print(f"\n🔍 Processing: {filename}")
        print(f"   📂 Collection: {collection_name} (ID: {collection_id})")
        print(f"   📖 Book number: {file_number}")
        
        # Find matching book in database
        book_match = find_matching_book(cursor, file_number, collection_id)
        if not book_match:
            print(f"❌ No book found in {collection_name} collection for number: {file_number}")
            not_found_count += 1
            continue
        
        book_id, db_number = book_match
        print(f"✅ Found matching book: {db_number} (ID: {book_id})")
        
        try:
            data = extract_data_from_docx(docx_path)
        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")
            continue

        if not data:
            print(f"⚠️ No metadata extracted from {filename}")
            continue

        # Insert or update description in book_descriptions table
        try:
            action = insert_or_update_description(cursor, book_id, collection_id, db_number, data)
            if action == "updated":
                print(f"✅ Updated description for book {db_number}")
                updated_count += 1
            else:
                print(f"✅ Inserted new description for book {db_number}")
                inserted_count += 1
            
            # Print some sample data to verify hyperlinks
            for col, value in list(data.items())[:2]:
                if '<a href=' in str(value):
                    print(f"   📎 {col}: {value[:100]}...")
        except sqlite3.Error as e:
            print(f"❌ Database error for book {db_number}: {e}")
        
        processed_count += 1

    conn.commit()
    conn.close()
    
    print(f"\n📊 Summary:")
    print(f"   📄 Files processed: {processed_count}")
    print(f"   ✅ Descriptions inserted: {inserted_count}")
    print(f"   🔄 Descriptions updated: {updated_count}")
    print(f"   ❌ Books not found: {not_found_count}")
    print(f"   📚 Total cinquecentine books: {cinquecentine_books}")
    print(f"   📚 Total incunaboli books: {incunaboli_books}")

if __name__ == "__main__":
    update_db_from_docx()